import json
import os
import re
import secrets
import shutil
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional, Set
from uuid import uuid4
from datetime import datetime

from cryptography.hazmat.primitives.kdf.scrypt import Scrypt
from docx import Document
from pypdf import PdfReader

DATA_DIR = Path.cwd().parent.parent / 'local-data'
UPLOADS_DIR = DATA_DIR / 'uploads'
DB_FILE = DATA_DIR / 'db.json'
ROOT_DIR = Path(__file__).resolve().parents[1]

if str(ROOT_DIR) not in sys.path:
    sys.path.append(str(ROOT_DIR))

try:
    from matcher.service import matcher_service  # type: ignore
except Exception:
    matcher_service = None

STOPWORDS: Set[str] = {
    'a', 'an', 'and', 'the', 'to', 'of', 'for', 'in', 'on', 'with', 'as', 'at', 'by', 'or', 'from', 'is', 'are', 'be',
    'this', 'that', 'these', 'those', 'will', 'can', 'must', 'should', 'have', 'has', 'had', 'our', 'your', 'their',
    'we', 'you', 'they', 'he', 'she', 'it', 'but', 'not', 'than', 'into', 'about', 'over', 'under', 'per', 'using',
    'used', 'across', 'between', 'within', 'ability', 'skills', 'skill', 'experience', 'work', 'job', 'role', 'resume',
    'candidate', 'requirements', 'preferred', 'plus', 'good', 'strong', 'basic', 'advanced', 'knowledge', 'understanding',
    'и', 'в', 'на', 'с', 'для', 'по', 'к', 'из', 'как', 'или', 'это', 'что', 'а', 'но', 'не', 'при', 'быть', 'есть',
    'будет', 'нужно', 'опыт', 'резюме', 'кандидат', 'требования', 'навыки', 'жұмыс', 'үшін', 'және', 'дағды', 'талаптар'
}

SKILL_ALIASES: Dict[str, List[str]] = {
    'python': ['python'],
    'java': ['java'],
    'javascript': ['javascript', 'js'],
    'typescript': ['typescript', 'ts'],
    'react': ['react', 'react.js', 'reactjs'],
    'next.js': ['next.js', 'nextjs'],
    'node.js': ['node.js', 'nodejs'],
    'fastapi': ['fastapi'],
    'django': ['django'],
    'flask': ['flask'],
    'sql': ['sql', 'postgresql', 'mysql', 'sqlite'],
    'mongodb': ['mongodb', 'mongo db', 'mongo'],
    'docker': ['docker'],
    'kubernetes': ['kubernetes', 'k8s'],
    'aws': ['aws', 'amazon web services'],
    'azure': ['azure'],
    'gcp': ['gcp', 'google cloud'],
    'git': ['git', 'github', 'gitlab'],
    'rest api': ['rest api', 'restful api', 'api development'],
    'graphql': ['graphql'],
    'html': ['html', 'html5'],
    'css': ['css', 'css3', 'scss', 'sass'],
    'tailwind css': ['tailwind', 'tailwindcss'],
    'machine learning': ['machine learning', 'ml'],
    'deep learning': ['deep learning'],
    'nlp': ['nlp', 'natural language processing'],
    'bert': ['bert', 'sbert', 'sentence-bert'],
    'pandas': ['pandas'],
    'numpy': ['numpy'],
    'power bi': ['power bi'],
    'tableau': ['tableau'],
    'excel': ['excel', 'microsoft excel'],
    'project management': ['project management'],
    'communication': ['communication', 'communicator'],
    'leadership': ['leadership', 'team lead'],
    'problem solving': ['problem solving', 'problem-solving'],
    'testing': ['testing', 'qa', 'quality assurance'],
    'linux': ['linux'],
    'c++': ['c++', 'cpp'],
    'c#': ['c#', 'csharp'],
    'matlab': ['matlab'],
    'scada': ['scada'],
    'plc': ['plc'],
    'automation': ['automation', 'industrial automation'],
}

EDUCATION_LEVELS = ['phd', 'master', 'bachelor', 'associate', 'diploma']


def default_db() -> Dict[str, List[Any]]:
    return {
        'users': [],
        'profiles': [],
        'user_settings': [],
        'resumes': [],
        'scan_results': [],
    }


def ensure_local_storage():
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    if not DB_FILE.exists():
        with open(DB_FILE, 'w', encoding='utf-8') as f:
            json.dump(default_db(), f, indent=2)


def read_db():
    ensure_local_storage()
    with open(DB_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def write_db(db):
    ensure_local_storage()
    with open(DB_FILE, 'w', encoding='utf-8') as f:
        json.dump(db, f, indent=2, ensure_ascii=False)


def hash_password(password: str) -> str:
    salt = secrets.token_bytes(16).hex()
    kdf = Scrypt(salt=bytes.fromhex(salt), length=64, n=2**14, r=8, p=1)
    derived = kdf.derive(password.encode())
    return f'{salt}:{derived.hex()}'


def verify_password(password: str, password_hash: str) -> bool:
    try:
        salt_hex, stored_hex = password_hash.split(':', 1)
        salt = bytes.fromhex(salt_hex)
        kdf = Scrypt(salt=salt, length=64, n=2**14, r=8, p=1)
        derived = kdf.derive(password.encode())
        stored = bytes.fromhex(stored_hex)
        return secrets.compare_digest(stored, derived)
    except Exception:
        return False


def to_client_user(user):
    return {
        'id': user['id'],
        'email': user['email'],
        'user_metadata': {
            'full_name': user['full_name'],
        },
    }


def create_user(email: str, password: str, full_name: str):
    normalized_email = email.strip().lower()
    db = read_db()
    if any(u['email'] == normalized_email for u in db['users']):
        raise ValueError('User with this email already exists')
    now = datetime.now().isoformat()
    uid = str(uuid4())
    pwd_hash = hash_password(password)
    user = {
        'id': uid,
        'email': normalized_email,
        'password_hash': pwd_hash,
        'full_name': full_name.strip(),
        'created_at': now,
    }
    profile = {
        'id': uid,
        'email': normalized_email,
        'full_name': full_name.strip(),
        'created_at': now,
        'updated_at': now,
    }
    settings = {
        'user_id': uid,
        'auto_save_resumes': True,
        'color_scheme': 'emerald',
        'language': 'en',
        'max_storage_mb': 500,
        'created_at': now,
        'updated_at': now,
    }
    db['users'].append(user)
    db['profiles'].append(profile)
    db['user_settings'].append(settings)
    write_db(db)
    return to_client_user(user)


def authenticate_user(email: str, password: str):
    normalized_email = email.strip().lower()
    db = read_db()
    user = next((u for u in db['users'] if u['email'] == normalized_email), None)
    if not user or not verify_password(password, user['password_hash']):
        raise ValueError('Invalid email or password')
    return to_client_user(user)


def get_profile(user_id: str):
    db = read_db()
    return next((p for p in db['profiles'] if p['id'] == user_id), None)


def get_settings(user_id: str):
    db = read_db()
    return next((s for s in db['user_settings'] if s['user_id'] == user_id), None)


def sanitize_name(name: str):
    return re.sub(r'[^a-zA-Z0-9._-]', '_', name)


def normalize_spaces(text: str) -> str:
    return re.sub(r'\s+', ' ', (text or '').replace('\x00', ' ')).strip()


def normalize_for_search(text: str) -> str:
    cleaned = normalize_spaces(text).lower()
    cleaned = cleaned.replace('/', ' ').replace('-', ' ')
    return cleaned


def clean_extracted_text(text: str) -> str:
    lines = [normalize_spaces(line) for line in (text or '').splitlines()]
    cleaned = '\n'.join(line for line in lines if line)
    return cleaned.strip()

def tokenize(text: str) -> List[str]:
    return [token for token in re.findall(r'[a-zA-Zа-яА-ЯёЁқҚңҢғҒүҮұҰөӨһҺ0-9\+#\.]+', normalize_for_search(text)) if token not in STOPWORDS and len(token) > 1]


def unique_preserve(items: List[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for item in items:
        if item and item not in seen:
            seen.add(item)
            out.append(item)
    return out


def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    chunks: List[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or '')
    return '\n'.join(chunks)


def extract_text_from_docx(file_path: str) -> str:
    document = Document(file_path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def extract_text_from_file(file_path: str, file_type: str, file_name: str) -> str:
    ext = Path(file_name).suffix.lower()
    try:
        if file_type == 'text/plain' or ext == '.txt':
            return Path(file_path).read_text(encoding='utf-8', errors='ignore')
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        if ext == '.docx':
            return extract_text_from_docx(file_path)
    except Exception:
        pass
    return f'[Preview text unavailable for {file_name}]'


def extract_email(text: str) -> Optional[str]:
    match = re.search(r'[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}', text, flags=re.IGNORECASE)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(r'(\+?\d[\d\s\-\(\)]{7,}\d)', text)
    return normalize_spaces(match.group(1)) if match else None


def extract_name(text: str) -> Optional[str]:
    lines = [normalize_spaces(line) for line in text.splitlines() if normalize_spaces(line)]
    for line in lines[:5]:
        if '@' in line or any(char.isdigit() for char in line) or len(line.split()) > 5:
            continue
        if re.fullmatch(r'[A-Za-zА-Яа-яЁёҚқҒғҮүҰұӨөҺһ\-\']+(\s+[A-Za-zА-Яа-яЁёҚқҒғҮүҰұӨөҺһ\-\']+){1,3}', line):
            return line
    return None


def extract_skills(text: str) -> List[str]:
    searchable = normalize_for_search(text)
    detected: List[str] = []
    for canonical, aliases in SKILL_ALIASES.items():
        for alias in aliases:
            normalized_alias = normalize_for_search(alias)
            pattern = r'(?<![a-zA-Zа-яА-ЯёЁқҚңҢғҒүҮұҰөӨһҺ0-9])' + re.escape(normalized_alias).replace('\\ ', r'\s+') + r'(?![a-zA-Zа-яА-ЯёЁқҚңҢғҒүҮұҰөӨһҺ0-9])'
            if re.search(pattern, searchable, flags=re.IGNORECASE):
                detected.append(canonical)
                break
    return unique_preserve(sorted(detected))


def extract_education_level(text: str) -> Optional[str]:
    lowered = normalize_for_search(text)
    patterns = {
        'phd': ['phd', 'doctor of philosophy', 'doctoral'],
        'master': ['master', 'msc', 'm.sc', 'магистр'],
        'bachelor': ['bachelor', 'bsc', 'b.sc', 'бакалавр'],
        'associate': ['associate degree'],
        'diploma': ['diploma', 'college diploma'],
    }
    for level in EDUCATION_LEVELS:
        for alias in patterns[level]:
            if alias in lowered:
                return level
    return None


def education_rank(level: Optional[str]) -> int:
    if level is None:
        return 0
    ranking = {'diploma': 1, 'associate': 2, 'bachelor': 3, 'master': 4, 'phd': 5}
    return ranking.get(level, 0)


def extract_years_of_experience(text: str) -> Optional[int]:
    lowered = normalize_for_search(text)
    patterns = [
        r'(\d{1,2})\+?\s*(?:years|year|yrs|yr)',
        r'(?:experience of|minimum of|at least)\s*(\d{1,2})\+?\s*(?:years|year|yrs|yr)',
    ]
    values: List[int] = []
    for pattern in patterns:
        for match in re.findall(pattern, lowered):
            try:
                values.append(int(match))
            except Exception:
                continue
    if values:
        return max(values)

    if 'senior' in lowered:
        return 5
    if 'mid-level' in lowered or 'mid level' in lowered:
        return 3
    if 'junior' in lowered or 'entry level' in lowered or 'entry-level' in lowered or 'intern' in lowered:
        return 1
    return None


def top_keywords_from_text(text: str, limit: int = 12) -> List[str]:
    tokens = tokenize(text)
    frequency: Dict[str, int] = {}
    for token in tokens:
        frequency[token] = frequency.get(token, 0) + 1
    ranked = sorted(frequency.items(), key=lambda item: (-item[1], item[0]))
    return [token for token, _ in ranked[:limit]]


def parse_resume_profile(text: str) -> Dict[str, Any]:
    skills = extract_skills(text)
    return {
        'name': extract_name(text),
        'email': extract_email(text),
        'phone': extract_phone(text),
        'skills': skills,
        'education_level': extract_education_level(text),
        'years_experience': extract_years_of_experience(text),
        'top_keywords': top_keywords_from_text(text),
    }


def parse_job_description(text: str) -> Dict[str, Any]:
    skills = extract_skills(text)
    return {
        'required_skills': skills,
        'education_level': extract_education_level(text),
        'years_experience': extract_years_of_experience(text),
        'top_keywords': top_keywords_from_text(text),
    }


def keyword_overlap_score(source_keywords: List[str], target_keywords: List[str]) -> int:
    if not target_keywords:
        return 100 if source_keywords else 0
    overlap = len(set(source_keywords) & set(target_keywords))
    return round((overlap / max(len(set(target_keywords)), 1)) * 100)


def context_similarity_score(text_a: str, text_b: str) -> int:
    tokens_a = set(tokenize(text_a))
    tokens_b = set(tokenize(text_b))
    if not tokens_a or not tokens_b:
        return 0
    intersection = len(tokens_a & tokens_b)
    union = len(tokens_a | tokens_b)
    return round((intersection / union) * 100)


def experience_alignment_score(resume_years: Optional[int], required_years: Optional[int]) -> int:
    if required_years is None:
        return 100 if resume_years is not None else 70
    if resume_years is None:
        return 35
    return min(100, round((resume_years / max(required_years, 1)) * 100))


def education_alignment_score(resume_level: Optional[str], required_level: Optional[str]) -> int:
    if required_level is None:
        return 100 if resume_level else 70
    resume_rank = education_rank(resume_level)
    required_rank = education_rank(required_level)
    if resume_rank == 0:
        return 40
    if resume_rank >= required_rank:
        return 100
    return max(45, 100 - ((required_rank - resume_rank) * 20))


def build_match_summary(matched_skills: List[str], missing_skills: List[str], resume_profile: Dict[str, Any], total_score: int) -> str:
    parts: List[str] = []
    if matched_skills:
        parts.append(f"Strong alignment in {', '.join(matched_skills[:4])}.")
    if missing_skills:
        parts.append(f"Missing or weak evidence for {', '.join(missing_skills[:4])}.")
    years = resume_profile.get('years_experience')
    education = resume_profile.get('education_level')
    if years is not None:
        parts.append(f'Estimated experience level: {years} year(s).')
    if education:
        parts.append(f'Detected education: {education.title()}.')
    if not parts:
        parts.append('The resume has limited structured signals, so the score is driven mainly by general context overlap.')
    parts.append(f'Overall job fit score: {total_score}%.')
    return ' '.join(parts)


def analyze_job_fit(resume_text: str, job_description: str) -> Dict[str, Any]:
    resume_profile = parse_resume_profile(resume_text)
    job_profile = parse_job_description(job_description)

    matched_skills = unique_preserve([skill for skill in resume_profile['skills'] if skill in job_profile['required_skills']])
    missing_skills = unique_preserve([skill for skill in job_profile['required_skills'] if skill not in resume_profile['skills']])

    skill_score = keyword_overlap_score(resume_profile['skills'], job_profile['required_skills'])
    experience_score = experience_alignment_score(resume_profile.get('years_experience'), job_profile.get('years_experience'))
    education_score = education_alignment_score(resume_profile.get('education_level'), job_profile.get('education_level'))
    keyword_score = keyword_overlap_score(resume_profile['top_keywords'], job_profile['top_keywords'])
    context_score = context_similarity_score(resume_text, job_description)

    total_score = round(
        (skill_score * 0.5) +
        (experience_score * 0.2) +
        (education_score * 0.1) +
        (keyword_score * 0.1) +
        (context_score * 0.1)
    )

    summary = build_match_summary(matched_skills, missing_skills, resume_profile, total_score)

    return {
        'match_score': total_score,
        'skill_score': skill_score,
        'experience_score': experience_score,
        'education_score': education_score,
        'context_score': context_score,
        'matched_skills': matched_skills,
        'missing_skills': missing_skills,
        'resume_profile': resume_profile,
        'job_profile': job_profile,
        'summary': summary,
    }


def save_resume_from_file(user_id: str, file_path: str, file_name: str, file_type: str, file_size: int):
    db = read_db()
    ext = Path(file_name).suffix or ''
    stored_name = f"{int(datetime.now().timestamp())}-{str(uuid4())}{ext}"
    user_dir = UPLOADS_DIR / user_id
    user_dir.mkdir(parents=True, exist_ok=True)
    stored_path = user_dir / sanitize_name(stored_name)
    shutil.copy2(file_path, stored_path)

    extracted_text = clean_extracted_text(extract_text_from_file(str(stored_path), file_type, file_name))
    resume_profile = parse_resume_profile(extracted_text)

    resume = {
        'id': str(uuid4()),
        'user_id': user_id,
        'file_name': file_name,
        'file_url': f'/api/files/{user_id}/{stored_name}',
        'file_type': file_type,
        'file_size': file_size,
        'extracted_text': extracted_text,
        'stored_name': stored_name,
        'resume_profile': resume_profile,
        'created_at': datetime.now().isoformat(),
    }
    db['resumes'].append(resume)
    write_db(db)
    return resume


def list_resumes(user_id: str):
    db = read_db()
    resumes = [r for r in db['resumes'] if r['user_id'] == user_id]
    results = []
    for resume in resumes:
        scans = [s for s in db['scan_results'] if s['resume_id'] == resume['id']]
        best_match_score = max([s['match_score'] for s in scans], default=-1) if scans else None
        results.append({
            **resume,
            'best_match_score': best_match_score,
            'resume_profile': resume.get('resume_profile') or parse_resume_profile(resume.get('extracted_text', '')),
        })
    results.sort(key=lambda x: datetime.fromisoformat(x['created_at']), reverse=True)
    return results


def delete_resume(user_id: str, resume_id: str):
    db = read_db()
    resume = next((r for r in db['resumes'] if r['id'] == resume_id and r['user_id'] == user_id), None)
    if not resume:
        raise ValueError('Resume not found')
    file_path = UPLOADS_DIR / user_id / resume['stored_name']
    if file_path.exists():
        file_path.unlink()
    db['resumes'] = [r for r in db['resumes'] if r['id'] != resume_id]
    db['scan_results'] = [s for s in db['scan_results'] if s['resume_id'] != resume_id]
    write_db(db)


def scan_resumes(user_id: str, keywords: Optional[List[str]] = None, job_description: Optional[str] = None, mode: str = 'keywords'):
    db = read_db()
    resumes = [r for r in db['resumes'] if r['user_id'] == user_id]
    if not resumes:
        return []

    batch = []
    best_resume_id = ''
    best_score = -1
    now = datetime.now().isoformat()

    if mode == 'job_description':
        jd_text = normalize_spaces(job_description or '')
        if not jd_text:
            raise ValueError('Job description is required')

        if matcher_service is not None:
            resume_records = [
                {
                    'id': resume['id'],
                    'text': resume.get('extracted_text', ''),
                    'file_name': resume.get('file_name', ''),
                    'file_type': resume.get('file_type', ''),
                    'resume_profile': resume.get('resume_profile', {}),
                }
                for resume in resumes
            ]

            ranked = matcher_service.rank(
                job_text=jd_text,
                resume_records=resume_records,
                top_k=5,
                explain=True,
            )

            best_resume_id = ranked[0]['resume_id'] if ranked else ''
            saved = []
            for entry in ranked:
                resume = next((r for r in resumes if r['id'] == entry['resume_id']), None)
                if resume is None:
                    continue

                score_pct = round(float(entry.get('score', 0.0)) * 100)
                skill_pct = round(float(entry.get('coverage', 0.0)) * 100)
                semantic_pct = round(float(entry.get('semantic', 0.0)) * 100)

                record = {
                    'id': str(uuid4()),
                    'user_id': user_id,
                    'resume_id': resume['id'],
                    'mode': 'job_description',
                    'job_description_excerpt': jd_text[:500],
                    'match_score': score_pct,
                    'matched_keywords': entry.get('matched', []),
                    'is_best_match': resume['id'] == best_resume_id,
                    'analysis': {
                        'summary': entry.get('explanation') or f"Overall suitability score: {score_pct}%.",
                        'missing_skills': entry.get('missing', []),
                        'skill_score': skill_pct,
                        'experience_score': 0,
                        'education_score': 0,
                        'context_score': semantic_pct,
                        'resume_profile': resume.get('resume_profile') or parse_resume_profile(resume.get('extracted_text', '')),
                    },
                    'created_at': now,
                }
                db['scan_results'].append(record)
                saved.append({
                    'id': record['id'],
                    'resumeId': resume['id'],
                    'fileName': resume['file_name'],
                    'fileType': resume['file_type'],
                    'matchScore': score_pct,
                    'matchedKeywords': entry.get('matched', []),
                    'missingSkills': entry.get('missing', []),
                    'summary': record['analysis']['summary'],
                    'skillScore': skill_pct,
                    'experienceScore': 0,
                    'educationScore': 0,
                    'contextScore': semantic_pct,
                    'detectedSkills': record['analysis']['resume_profile'].get('skills', []),
                    'educationLevel': record['analysis']['resume_profile'].get('education_level'),
                    'yearsExperience': record['analysis']['resume_profile'].get('years_experience'),
                    'isBestMatch': record['is_best_match'],
                })

            write_db(db)
            saved.sort(key=lambda x: x['matchScore'], reverse=True)
            return saved

        for resume in resumes:
            analysis = analyze_job_fit(resume.get('extracted_text', ''), jd_text)
            score = analysis['match_score']
            if score > best_score:
                best_score = score
                best_resume_id = resume['id']
            batch.append({'resume': resume, 'analysis': analysis, 'score': score})

        saved = []
        for item in batch:
            analysis = item['analysis']
            record = {
                'id': str(uuid4()),
                'user_id': user_id,
                'resume_id': item['resume']['id'],
                'mode': 'job_description',
                'job_description_excerpt': jd_text[:500],
                'match_score': item['score'],
                'matched_keywords': analysis['matched_skills'],
                'is_best_match': item['resume']['id'] == best_resume_id,
                'analysis': analysis,
                'created_at': now,
            }
            db['scan_results'].append(record)
            saved.append({
                'id': record['id'],
                'resumeId': item['resume']['id'],
                'fileName': item['resume']['file_name'],
                'fileType': item['resume']['file_type'],
                'matchScore': item['score'],
                'matchedKeywords': analysis['matched_skills'],
                'missingSkills': analysis['missing_skills'],
                'summary': analysis['summary'],
                'skillScore': analysis['skill_score'],
                'experienceScore': analysis['experience_score'],
                'educationScore': analysis['education_score'],
                'contextScore': analysis['context_score'],
                'detectedSkills': analysis['resume_profile']['skills'],
                'educationLevel': analysis['resume_profile'].get('education_level'),
                'yearsExperience': analysis['resume_profile'].get('years_experience'),
                'isBestMatch': record['is_best_match'],
            })
        write_db(db)
        saved.sort(key=lambda x: x['matchScore'], reverse=True)
        return saved

    normalized_keywords = [k.strip() for k in (keywords or []) if k.strip()]
    if not normalized_keywords:
        raise ValueError('At least one keyword required')

    for resume in resumes:
        text = normalize_for_search(resume.get('extracted_text', ''))
        matched_keywords = [k for k in normalized_keywords if normalize_for_search(k) in text]
        score = round((len(matched_keywords) / len(normalized_keywords)) * 100) if normalized_keywords else 0
        if score > best_score:
            best_score = score
            best_resume_id = resume['id']
        batch.append({'resume': resume, 'matched_keywords': matched_keywords, 'score': score})

    saved = []
    for item in batch:
        record = {
            'id': str(uuid4()),
            'user_id': user_id,
            'resume_id': item['resume']['id'],
            'mode': 'keywords',
            'keywords': normalized_keywords,
            'match_score': item['score'],
            'matched_keywords': item['matched_keywords'],
            'is_best_match': item['resume']['id'] == best_resume_id,
            'created_at': now,
        }
        db['scan_results'].append(record)
        saved.append({
            'id': record['id'],
            'resumeId': item['resume']['id'],
            'fileName': item['resume']['file_name'],
            'fileType': item['resume']['file_type'],
            'matchScore': item['score'],
            'matchedKeywords': item['matched_keywords'],
            'isBestMatch': record['is_best_match'],
        })
    write_db(db)
    saved.sort(key=lambda x: x['matchScore'], reverse=True)
    return saved


def get_dashboard(user_id: str):
    db = read_db()
    resumes = [r for r in db['resumes'] if r['user_id'] == user_id]
    scans = [s for s in db['scan_results'] if s['user_id'] == user_id]
    recent_scans = sorted(scans, key=lambda s: datetime.fromisoformat(s['created_at']), reverse=True)[:5]
    recent_scans = [
        {
            **s,
            'resumes': next(({'id': r['id'], 'file_name': r['file_name'], 'created_at': r['created_at']} for r in db['resumes'] if r['id'] == s['resume_id']), None)
        }
        for s in recent_scans
    ]
    return {
        'stats': {
            'resumesScanned': len(resumes),
            'keywordsMatched': len(scans),
            'bestMatches': len([s for s in scans if s['is_best_match']]),
        },
        'recentScans': recent_scans,
    }


def read_stored_file(slug: List[str]):
    file_path = UPLOADS_DIR / Path(*slug)
    if not file_path.exists():
        raise ValueError('File not found')
    return file_path.read_bytes()


def save_profile_and_settings(user_id: str, profile_input: Dict[str, Any], settings_input: Dict[str, Any]):
    db = read_db()
    now = datetime.now().isoformat()
    profile = next((p for p in db['profiles'] if p['id'] == user_id), None)
    if profile:
        profile['full_name'] = profile_input.get('full_name', profile['full_name'])
        profile['updated_at'] = now
    else:
        profile = {
            'id': user_id,
            'email': profile_input.get('email', ''),
            'full_name': profile_input.get('full_name', ''),
            'created_at': now,
            'updated_at': now,
        }
        db['profiles'].append(profile)

    user = next((u for u in db['users'] if u['id'] == user_id), None)
    if user and 'full_name' in profile_input:
        user['full_name'] = profile_input['full_name']

    settings = next((s for s in db['user_settings'] if s['user_id'] == user_id), None)
    if settings:
        settings.update({
            'auto_save_resumes': settings_input.get('auto_save_resumes', settings['auto_save_resumes']),
            'color_scheme': settings_input.get('color_scheme', settings['color_scheme']),
            'language': settings_input.get('language', settings['language']),
            'max_storage_mb': settings_input.get('max_storage_mb', settings['max_storage_mb']),
            'updated_at': now,
        })
    else:
        settings = {
            'user_id': user_id,
            'auto_save_resumes': settings_input.get('auto_save_resumes', True),
            'color_scheme': settings_input.get('color_scheme', 'emerald'),
            'language': settings_input.get('language', 'en'),
            'max_storage_mb': settings_input.get('max_storage_mb', 500),
            'created_at': now,
            'updated_at': now,
        }
        db['user_settings'].append(settings)

    write_db(db)
    return get_profile(user_id), get_settings(user_id)


def update_account(user_id: str, input_data: Dict[str, Any]):
    db = read_db()
    user_idx = next((i for i, u in enumerate(db['users']) if u['id'] == user_id), -1)
    if user_idx == -1:
        raise ValueError('User not found')
    user = db['users'][user_idx]
    profile_idx = next((i for i, p in enumerate(db['profiles']) if p['id'] == user_id), -1)

    next_email = (input_data.get('email') or user['email']).strip().lower()
    next_full_name = (input_data.get('full_name') or user['full_name']).strip()
    wants_email_change = next_email != user['email']
    wants_pwd_change = bool(input_data.get('new_password'))

    if wants_email_change or wants_pwd_change:
        if not input_data.get('current_password'):
            raise ValueError('Current password required for changes')
        if not verify_password(input_data['current_password'], user['password_hash']):
            raise ValueError('Current password incorrect')

    if wants_email_change and any(u['email'] == next_email for i, u in enumerate(db['users']) if i != user_idx):
        raise ValueError('Email already exists')
    if wants_pwd_change and len(input_data['new_password']) < 6:
        raise ValueError('Password min 6 chars')

    user['email'] = next_email
    user['full_name'] = next_full_name
    if wants_pwd_change:
        user['password_hash'] = hash_password(input_data['new_password'])

    if profile_idx != -1:
        db['profiles'][profile_idx].update({
            'email': next_email,
            'full_name': next_full_name,
            'updated_at': datetime.now().isoformat(),
        })

    write_db(db)
    return get_profile(user_id), to_client_user(user)
